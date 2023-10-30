import telebot
from telebot.types import InlineKeyboardButton, InlineKeyboardMarkup
import os
import time
import logging
import json
import traceback
import docx



#the telegram api key will be retrieved from .env file
KEY = os.environ.get("TELETOKEN")
PASSWORD = os.environ.get("CONFIG_PASSWORD")

#load users info from json file
try:
    with open('user_data.json', 'r') as json_file:
        loaded_user_data = json.load(json_file)
    user_data = {int(key): value for key, value in loaded_user_data.items()}
except FileNotFoundError:
        user_data = {}

#load contacts info
try:
    with open('contacts.json', 'r') as json_file:
        contacts = json.load(json_file)
    contacts = {key: int(value) for key, value in contacts.items()}
except FileNotFoundError:
        contacts = {}


#get contact ID by title
def getID(recipient):
    return contacts.get(recipient)


#initialize bot
bot = telebot.TeleBot(KEY)

logging.basicConfig(format='%(asctime)s - %(name)s - %(levelname)s - %(message)s', level=logging.INFO)
logging.info('App is running.')

#required command list
commands = {
    'start'       : 'Start your request',
    'help'        : 'Gives you information about the available commands'
}



#list of buttons to display on the menu and their callback data as a dictionary
optionset = {
    'IT dəstək': 'IT dəstək',
    'Təşkilati dəstək': 'Təşkilati dəstək',
    'Görüşlər/Qəbul': 'Görüşlər/Qəbul',
    'Mənə məlumat lazımdır': 'Mənə məlumat lazımdır',
    'Təklif | Şikayətim var':'Təklif və ya Şikayətim var'
}

suggestion = {
    'Təklif':'Təklif',
    'Şikayət':'Şikayət'

}

form = {'Online': 'Online',
        'Fiziki': 'Fiziki'}

technical = {
    'Zoom meeting yaradılmalıdır': 'Zoom meeting yaradılmalıdır',
    'Katric yenilənməlidir': 'Katric yenilənməlidir',
    'Docnetə qoşulmur': 'Docnetə qoşulmur',
    'Outlook': 'Outlook',
    'İnternetə qoşulmur': 'İnternetə qoşulmur',
    'Digər': 'Digər IT dəstək'
}

support = {
    'Ofisdə çatışmır': 'Ofisdə çatışmır',
    'Ərizə nümünələri': 'Ərizə nümünələri'
}

needed = {
    'Su': 'Su',
    'Salfetka': 'Salfetka',
    'Təmizlik vasitələri': 'Təmizlik vasitələri',
    'Digər': 'Digər çatışmır'
}

documents = {
    'İşdən azad edilmə': 'Ərizə azad ed',
    'Vakant yerə həvalə': 'Ərizə həvalə-boş vəzifə',
    'Ödənişsiz məzuniyyət': 'Ərizə-ödənişsiz',
    'İstifadə olunmayan məzuniyyət günlərinin keçirilməsi': 'Ərizə keçirilmə -',
    'Əmək məzuniyyətindən geri çağırılma': 'Ərizə-geriçağırılma və əvəzgün',
    'Vəzifəyə keçirilmə': 'Ərizə keçirilmə -',
    'İşə qəbul': 'Ərizə işə qəbul',
    'Qalıq günlər hesabına məzuniyyət': 'Ərizə-qalıq',
    'Sadə əmək məzuniyyəti': 'Ərizə-sadə',
    'Digər': 'Digər ərizə'

}

meetings = {
    'Rəhbərlik': 'Rəhbərlik meeting',
    'RİŞ': 'RİŞ meeting',
    'ÜŞ': 'ÜŞ meeting',
    'LİOŞ': 'LİOŞ meeting',
    'TMŞ': 'TMŞ meeting'
}

info = {
    'PR': 'PR info',
    'RİŞ': 'RİŞ info',
    'ÜŞ': 'ÜŞ info',
    'LİOŞ': 'LİOŞ info',
    'TMŞ': 'TMŞ info'

}

back = {'Əsas menyuya qayıdın' : 'Əsas menyuya qayıdın'}

confirm_meeting = {'Təsdiq edin': 'Görüşünüz təsdiqləndi',
                   'Imtina' : 'Seçdiyiniz vaxt uyğun deyil, yeni tarix və saat daxil edin.'}

confirm = {'✅': 'Müraciətiniz qeydə alındı'}

link = {'Görüşün link və detallarını göndərin' : 'Linki və təfərrüatları daxil edin: '}

reply = {'Cavablandırın': 'Cavablandırın'}






request_recipient = {}
feedback_destination = {}
zoom_meeting = {}
attempts = {}
start_times = {}


def set_destination(destination, user_id):

    request_recipient[user_id] = destination


def get_template(file_name):

    file_path = f"{file_name}.docx"

    try:
        doc = docx.Document(file_path)
        content = []
        for paragraph in doc.paragraphs:
            content.append(paragraph.text)
        return "\n".join(content)
    except FileNotFoundError:
        return f"File '{file_name}.docx' not found."


def custom_tech_handler(message):

        bot.send_message(chat_id = getID("IT dəstək"),text = f"{user_data[message.chat.id]}: {message.text}", reply_markup = gen_markup(confirm))
        feedback_destination[getID("IT dəstək")] = message.chat.id

        bot.send_message(message.chat.id, "Mesajınız göndərildi.", reply_markup = gen_markup(optionset))

def custom_request_handler(message):

        bot.send_message(chat_id = getID("Ofisdə çatışmır"),text = f"{user_data[message.chat.id]}: {message.text}", reply_markup = gen_markup(confirm))
        feedback_destination[getID("Ofisdə çatışmır")] = message.chat.id

        bot.send_message(message.chat.id, "Mesajınız göndərildi.", reply_markup = gen_markup(optionset))


def custom_text_handler(message):

        bot.send_message(chat_id = getID("Ofisdə çatışmır"),text = f"{user_data[message.chat.id]}: {message.text}", reply_markup = gen_markup(confirm))
        feedback_destination[getID("Ofisdə çatışmır")] = message.chat.id
        bot.send_message(message.chat.id, "Mesajınız göndərildi.", reply_markup = gen_markup(optionset))


def countdown_remaining_time(user_id):
    if user_id in start_times:
        elapsed_time = time.time() - start_times[user_id]
        remaining_time_seconds = max(0, 300 - elapsed_time)  # 300 seconds = 5 minutes
        remaining_time_minutes = remaining_time_seconds / 60
        return remaining_time_minutes
    else:
        return 0

def get_template(file_name):

    file_path = f"templates/{file_name}.docx"

    try:
        doc = docx.Document(file_path)
        content = []
        for paragraph in doc.paragraphs:
            content.append(paragraph.text)
        return "\n".join(content)
    except FileNotFoundError:
        return f"File '{file_name}.docx' not found."



def password_handler(message):
    user_id = message.from_user.id

    if message.text == PASSWORD and user_id not in start_times:
        bot.send_message(chat_id = user_id, text = "Adınızı və soyadınızı daxil edin: ")
        bot.register_next_step_handler(message, name_handler)
    else:
        if user_id not in attempts:
            attempts[user_id] = 1
        else:
            attempts[user_id] += 1

        if attempts[user_id] >= 3:
            if user_id not in start_times:
                start_times[user_id] = time.time()

            remaining_minutes = countdown_remaining_time(user_id)

            if remaining_minutes > 0:
                bot.send_message(chat_id = user_id, text = f"Növbəti cəhddən əvvəl {int(remaining_minutes) + 1} dəqiqə gözləyin.")
                bot.register_next_step_handler(message, password_handler)
            else:
                attempts.pop(user_id, None)
                start_times.pop(user_id, None)
                bot.register_next_step_handler(message, password_handler)
        else:
            bot.send_message(chat_id = user_id, text = "Təkrar cəhd edin.")
            bot.register_next_step_handler(message, password_handler)



def gen_markup(data_dict):
    markup = InlineKeyboardMarkup()
    buttons = [InlineKeyboardButton(text=text, callback_data=data) for text, data in data_dict.items()]
    markup.add(*buttons)
    return markup


def info_handler(message):

        bot.send_message(chat_id = getID(recipient=request_recipient[message.chat.id]),text = f"{user_data[message.chat.id]} soruşur: {message.text}", reply_markup = gen_markup(reply))
        bot.send_message(chat_id = message.chat.id, text = "Mesajınız göndərildi.", reply_markup = gen_markup(optionset))


def suggestion_handler(message):

        bot.send_message(chat_id = getID('Təklif'),text = f"{user_data[message.chat.id]}\n Təklif: {message.text}")
        bot.send_message(chat_id = message.chat.id, text = "Mesajınız göndərildi.", reply_markup = gen_markup(optionset))



def complaint_handler(message):

        bot.send_message(chat_id = getID('Şikayət'),text = f"{user_data[message.chat.id]}\n Şikayət: {message.text}")
        bot.send_message(chat_id = message.chat.id, text = "Mesajınız göndərildi.", reply_markup = gen_markup(optionset))




def  link_handler(message):
     user_id = message.chat.id
     bot.send_message(chat_id = feedback_destination[user_id], text = f"Zoom meeting link: {message.text}")
     


def zoom_topic_handler(message):

        user_id = message.chat.id
        if user_id not in zoom_meeting:
            zoom_meeting[user_id] = {}
        zoom_meeting[user_id]['topic'] = message.text

        bot.send_message(user_id, "Tarix daxil edin: ")
        bot.register_next_step_handler(message, zoom_date_handler)

def zoom_date_handler(message):
        user_id = message.chat.id
   
        if user_id in zoom_meeting:
            zoom_meeting[user_id]['date'] = message.text
            bot.send_message(user_id, "Saat daxil edin: ")
            bot.register_next_step_handler(message, zoom_time_handler)
        else:
            bot.send_message(user_id, "Zoom görüş datası düzgün formatda deyil.", reply_markup = gen_markup(optionset))


def zoom_time_handler(message):
        user_id = message.chat.id

        if user_id in zoom_meeting:

            zoom_meeting[user_id]['time'] = message.text

            user_info = zoom_meeting[user_id]
            info_string = f"Mövzu: {user_info['topic']}\nTarix: {user_info['date']}\nSaat: {user_info['time']}"



            bot.send_message(chat_id = getID("IT dəstək"),text = f"{user_data[user_id]} \nZoom meeting:\n{info_string}", reply_markup=gen_markup(link))
            feedback_destination[getID("IT dəstək")] = user_id
            bot.send_message(user_id, "Mesajınız göndərildi.", reply_markup=gen_markup(optionset))


            del zoom_meeting[user_id]
        else:
            bot.send_message(user_id, "Zoom görüş datası düzgün formatda deyil.", reply_markup = gen_markup(optionset))




def answer_handler(message):

    bot.send_message(chat_id = feedback_destination[message.chat.id], text = f"Sualınız cavablandırıldı: {message.text}")



def meeting_topic_handler(message):

        user_id = message.chat.id
        if user_id in meetings:
            meetings[user_id]['topic'] = message.text
            bot.send_message(user_id, "Tarix daxil edin: ")
            bot.register_next_step_handler(message, date_handler)




def date_handler(message):
    user_id = message.chat.id

    if user_id in meetings:
        meetings[user_id]['date'] = message.text
        bot.send_message(user_id, "Saat daxil edin: ")
        bot.register_next_step_handler(message, time_handler)
    else:
        bot.send_message(user_id, "Data düzgün formatda deyil.", reply_markup=gen_markup(optionset))




def time_handler(message):
    user_id = message.chat.id

    if user_id in meetings:
        meetings[user_id]['time'] = message.text
        if  meetings[user_id]['type'] == 'Fiziki':
            bot.send_message(user_id, "Məkan daxil edin: ")
            bot.register_next_step_handler(message, location_handler)
        elif  meetings[user_id]['type'] == 'Online':
            bot.send_message(user_id, "İştirakçıları daxil edin: ")
            bot.register_next_step_handler(message, attendees_handler)

    else:
        bot.send_message(user_id, "Data düzgün formatda deyil.", reply_markup=gen_markup(optionset))




def location_handler(message):
    user_id = message.chat.id

    if user_id in meetings:
        meetings[user_id]['location'] = message.text
        bot.send_message(user_id, "İştirakçıları daxil edin: ")
        bot.register_next_step_handler(message, attendees_handler)
    else:
        bot.send_message(user_id, "Data düzgün formatda deyil.", reply_markup=gen_markup(optionset))




def attendees_handler(message):
    user_id = message.chat.id

    if user_id in meetings:
        meetings[user_id]['attendees'] = message.text

        user_info = meetings[user_id]
        if  meetings[user_id]['type'] == 'Fiziki':
            info_string = f"{user_data[user_id]} \nTip: {user_info['type']}\n Mövzu: {user_info['topic']}\nTarix: {user_info['date']}\nSaat: {user_info['time']} \nMəkan: {user_info['location']} \nİştirakçılar: {user_info['attendees']}"


        elif  meetings[user_id]['type'] == 'Online':
            info_string = f"{user_data[user_id]} \nTip: {user_info['type']}\n Mövzu: {user_info['topic']}\nTarix: {user_info['date']}\nSaat: {user_info['time']} \nİştirakçılar: {user_info['attendees']}"

        bot.send_message(chat_id = getID(request_recipient[user_id]),text = f"Meeting:\n{info_string}", reply_markup=gen_markup(confirm_meeting))
        feedback_destination[getID(request_recipient[user_id])] = user_id
        bot.send_message(user_id, "Mesajınız göndərildi.", reply_markup=gen_markup(optionset))

        del meetings[user_id]
    else:
        bot.send_message(user_id, "Data düzgün formatda deyil.", reply_markup=gen_markup(optionset))
        del meetings[user_id]
        
        
        
        
        
        
        
        
@bot.message_handler(commands=['start'])
def command_start(message):
    logging.info(f"User {message.chat.id} started the bot.")
    user_id = message.from_user.id


    if user_id not in user_data:
        logging.info(f"New user detected - User ID: {user_id}")
        bot.send_message(chat_id =user_id, text = "Qeydiyyatdan keçmək üçün doğrulama parolunu daxil edin: ")
        bot.register_next_step_handler(message, password_handler)

    else:
        logging.info(f"User {message.from_user.id} is already registered")
        bot.send_message(chat_id =user_id, text = "Xoş gəldiniz.")
        bot.send_message(chat_id =user_id, text = f"{(user_data[message.from_user.id].split())[0]}, necə kömək edə bilərəm? Menyudan seçim edin.", reply_markup = gen_markup(optionset))




@bot.message_handler(commands=['help'])
def command_help(message):
    logging.info(f"User {message.from_user.id} requested help")
    chat_id = message.chat.id
    help_text = "Commands: \n"
    for key in commands:  
        help_text += "/" + key + ": "
        help_text += commands[key] + "\n"
    bot.send_message(chat_id, help_text) 



def save_user_data(user_id, message):
    user_name = message.text

    logging.info(f"does the id exist: {user_id not in user_data}")
    logging.info(f"user_data: {user_data}, user_id: {user_id}")

    if user_data.get(user_id, None) == None:

        user_data[user_id] = user_name

        with open('user_data.json', 'w') as json_file:
            json.dump(user_data, json_file)

        logging.info(f"New user detected - User ID: {user_id}, User Name: {user_name}")



def name_handler(message):
     user_id = message.from_user.id
     logging.info(f"User {message.from_user.id} entered his/her name.")
     save_user_data(message.chat.id, message)

     bot.send_message(chat_id = message.chat.id, text = f"{(user_data[user_id].split())[0]}, necə kömək edə bilərəm? Menyudan seçim edin.", reply_markup = gen_markup(optionset))


@bot.callback_query_handler(func=lambda call: True)
def callback_worker(call):
    '''Callback Handler:
    This function manages user interactions with the bot's inline keyboard buttons, ensuring appropriate responses.

    Steps:
    1. Extract user's ID and callback data from the interaction.
    2. Analyze the callback data to understand the user's intention.
    3. Trigger specific actions or responses based on the user's choice.
    '''
    user_id = call.from_user.id
    data = call.data

    logging.info(f"{user_data[user_id]}: triggered callback - {data}")
    
    
    if data == 'Digər IT dəstək':
            bot.send_message(chat_id = user_id, text ='Mesajınızı daxil edin: ')
            bot.register_next_step_handler(call.message, custom_tech_handler)


    elif data == 'Digər çatışmır':
            bot.send_message(chat_id = user_id, text ='Mesajınızı daxil edin: ')
            bot.register_next_step_handler(call.message, custom_request_handler)


    elif data == 'Digər ərizə':
            bot.send_message(chat_id = user_id, text ='Mesajınızı daxil edin: ')
            bot.register_next_step_handler(call.message, custom_text_handler)


    elif data == 'IT dəstək':
        bot.send_message(chat_id = user_id,text = "Menyudan seçim edin.", reply_markup = gen_markup(technical))


    elif data == 'Təşkilati dəstək' or data == 'Təşkilati dəstək':
        bot.send_message(chat_id = user_id, text = "Menyudan seçim edin.", reply_markup =  gen_markup(support))


    elif data == 'Görüşlər/Qəbul':
            bot.send_message(chat_id = user_id,text = "Menyudan seçim edin.", reply_markup =  gen_markup(meetings))


    elif data == 'Mənə məlumat lazımdır':
            bot.send_message(chat_id = user_id, text = "Menyudan seçim edin.", reply_markup =  gen_markup(info))


    elif data == 'Zoom meeting yaradılmalıdır':
            feedback_destination[getID("IT dəstək")] = user_id
            bot.send_message(chat_id = user_id, text = "Mövzu daxil edin: ")
            bot.register_next_step_handler(call.message, zoom_topic_handler)
            

    elif data in technical.values() and (data != 'Digər IT dəstək' and data != "Zoom meeting yaradılmalıdır"):
         bot.send_message(chat_id = getID("IT dəstək"), text = f"{user_data[user_id]}: {data}", reply_markup = gen_markup(confirm))
         feedback_destination[getID("IT dəstək")] = user_id
         bot.send_message(chat_id = user_id, text ="Mesajınız göndərildi.", reply_markup = gen_markup(optionset))
         

    elif len(data.split(" ")) >= 2 and data.split(" ")[1] == 'meeting':
        set_destination(data.split(" ")[0], user_id)
        bot.send_message(chat_id = user_id,text = "Görüşün növünü seçin: ", reply_markup = gen_markup(form))
        

    elif data == 'Online' or data == 'Fiziki':
        if user_id not in meetings:
                meetings[user_id] = {}
        meetings[user_id]['type'] = data
        bot.send_message(user_id, "Mövzu daxil edin: ")
        bot.register_next_step_handler(call.message, meeting_topic_handler)
        

    elif data == 'Ofisdə çatışmır':
           bot.send_message(chat_id = user_id,text = "Xahiş olunur menyudan seçim edin.", reply_markup = gen_markup(needed))
           

    elif data == 'Təklif və ya Şikayətim var':
         bot.send_message(chat_id = user_id,text = "Xahiş olunur menyudan seçim edin.", reply_markup = gen_markup(suggestion))
         

    elif data == 'Şikayət':
            bot.send_message(chat_id = user_id, text = "Şikayətinizi daxil edin: ")
            bot.register_next_step_handler(call.message, complaint_handler)
            

    elif data == 'Təklif':
            bot.send_message(chat_id = user_id, text = "Təklifinizi daxil edin: ")
            bot.register_next_step_handler(call.message, suggestion_handler)
            

    elif data == 'Ərizə nümünələri':
           bot.send_message(chat_id = user_id,text = "Xahiş olunur menyudan seçim edin.", reply_markup = gen_markup(documents))
           

    elif data in documents.values() and data != 'Digər ərizə':
            bot.send_message(chat_id = user_id, text = str(get_template(data)), reply_markup = gen_markup(optionset))
            

    elif data in needed.values() and data != 'Digər çatışmır':
        bot.send_message(chat_id =user_id, text = "Mesajınız göndərildi.", reply_markup = gen_markup(optionset))
        bot.send_message(chat_id = getID("Ofisdə çatışmır"), text = f"{user_data[user_id]}: {data}", reply_markup = gen_markup(confirm))
        feedback_destination[getID("Ofisdə çatışmır")] = user_id
        

    elif len(data.split(" ")) >= 2 and data.split(" ")[1] == 'info':

        set_destination((data.split(" "))[0], user_id)
        logging.info(request_recipient[user_id])
        feedback_destination[getID(request_recipient[user_id])] = user_id
        bot.send_message(chat_id = user_id , text = "Sorğu etmək istədiyiniz məlumatı daxil edin: ")
        bot.register_next_step_handler(call.message, info_handler)


    elif data == 'Cavablandırın':
        bot.send_message(chat_id = user_id , text = "Cavabınızı daxil edin: ")
        bot.register_next_step_handler(call.message, answer_handler)


    elif data == 'Görüşünüz təsdiqləndi.':
         bot.send_message(chat_id = feedback_destination[user_id], text = data)
         

    elif data == 'Müraciətiniz qeydə alındı.':
         bot.send_message(chat_id = feedback_destination[user_id], text = data)
         

    elif data == 'Seçdiyiniz vaxt uyğun deyil, yeni görüş yaradın.':
         bot.send_message(chat_id = feedback_destination[user_id], text = data)
         

    elif data == 'Linki və təfərrüatları daxil edin: ':
        bot.send_message(chat_id = user_id, text = data)
        bot.register_next_step_handler(call.message, link_handler)
       



if __name__ == "__main__":
    while True:
        try:
            bot.polling()
        except Exception as e:
            logging.error(f"Bot encountered an error: {e}\nTraceback: {traceback.format_exc()}", exc_info=True)
            # Add a delay before restarting to avoid spamming the API
            time.sleep(3)  # You can adjust the delay duration as needed